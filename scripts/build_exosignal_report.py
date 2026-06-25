from __future__ import annotations

import json
import math
import os
import textwrap
import urllib.request
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "deliverables"
ASSETS = OUT / "assets"
OUT.mkdir(exist_ok=True)
ASSETS.mkdir(exist_ok=True)

REPORT = OUT / "Rapport_ExoSignal_Observatory.docx"

SCREENSHOTS = {
    "analyze": Path("/tmp/exosignal-report-01-analyze.png"),
    "universe": Path("/tmp/exosignal-report-02-universe.png"),
    "planet": Path("/tmp/exosignal-report-03-planet.png"),
    "signals": Path("/tmp/exosignal-report-04-signals.png"),
    "predict": Path("/tmp/exosignal-report-05-predict.png"),
    "command": Path("/tmp/exosignal-report-06-command.png"),
}


def fetch_json(url: str, fallback):
    try:
        with urllib.request.urlopen(url, timeout=8) as response:
            return json.loads(response.read().decode("utf-8"))
    except Exception:
        return fallback


DATASET = fetch_json(
    "http://127.0.0.1:8000/api/dataset",
    {
        "repo": "bingbangboom/exoplanet-transit-detection",
        "full_dataset_rows": 23567,
        "missions": ["Kepler", "K2", "TESS"],
        "classes": ["PLANET", "FALSE_POSITIVE", "NO_SIGNAL"],
        "files": {
            "metadata": {"size_mb": 31.81, "path": "data/metadata.csv"},
            "train": {"size_mb": 1457.82, "path": "data/train.parquet"},
            "val": {"size_mb": 182.29, "path": "data/val.parquet"},
            "test": {"size_mb": 182.11, "path": "data/test.parquet"},
        },
        "eda": {
            "class_distribution": {"PLANET": 9791, "NO_SIGNAL": 7840, "FALSE_POSITIVE": 5936},
            "mission_distribution": {"kepler": 16054, "tess": 6028, "k2": 1485},
            "numeric_ranges": {
                "period_days": {"median": 5.60402307},
                "depth_ppm": {"median": 1310.6803156},
                "duration_hrs": {"median": 3.218},
                "planet_radius_earth": {"median": 5.09878},
            },
        },
    },
)

MODEL = fetch_json(
    "http://127.0.0.1:8000/api/model",
    {
        "trained": True,
        "metrics": {
            "model_name": "XGBoost baseline",
            "rows_used": 6000,
            "features": 58,
            "positive_rate": 0.41383333333333333,
            "precision": 0.8126009693053312,
            "recall": 0.9212454212454212,
            "f1": 0.863519313304721,
            "roc_auc": 0.9508026426631077,
            "pr_auc": 0.9138196045666765,
            "confusion_matrix": {"labels": ["not_planet", "planet"], "values": [[658, 116], [43, 503]]},
            "elapsed_seconds": 148.59,
        },
    },
)

DEMO = fetch_json(
    "http://127.0.0.1:8000/api/demo",
    {
        "features": {
            "point_count": 2400,
            "depth_estimate": 0.013091142141465895,
            "dip_count": 6,
            "period_estimate": 4.318466027511462,
            "duration_estimate": 0.40396801028836393,
            "snr_estimate": 4.745036956089203,
            "variability": 0.002231704184772699,
        },
        "prediction": {"candidate_probability": 0.7406337770786824},
    },
)


NAVY = RGBColor(11, 31, 58)
BLUE = RGBColor(11, 61, 145)
RED = RGBColor(252, 61, 33)
GRAY = RGBColor(82, 94, 111)
LIGHT = RGBColor(242, 246, 251)


def font(size=11, bold=False, color=None, name="Aptos"):
    def apply(run):
        run.font.name = name
        run.font.size = Pt(size)
        run.font.bold = bold
        if color:
            run.font.color.rgb = color
    return apply


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color=None, size=9.5):
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(str(text))
    font(size=size, bold=bold, color=color)(r)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    font(9, color=GRAY)(run)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def style_document(doc: Document):
    section = doc.sections[0]
    section.top_margin = Cm(2.1)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.1)
    section.right_margin = Cm(2.1)
    section.header_distance = Cm(1.1)
    section.footer_distance = Cm(1.1)

    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10.5)
    styles["Normal"].paragraph_format.space_after = Pt(6)
    styles["Normal"].paragraph_format.line_spacing = 1.12

    for name, size, color, before, after in [
        ("Heading 1", 17, BLUE, 14, 8),
        ("Heading 2", 13, NAVY, 10, 5),
        ("Heading 3", 11.5, GRAY, 7, 3),
    ]:
        st = styles[name]
        st.font.name = "Aptos Display"
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = color
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)


def add_footer(doc):
    for section in doc.sections:
        footer = section.footer
        p = footer.paragraphs[0]
        p.text = ""
        r = p.add_run("ExoSignal Observatory — Rapport de projet")
        font(8.5, color=GRAY)(r)
        p.add_run("   ")
        add_page_number(p)


def add_hline(paragraph, color="FC3D21", size="8"):
    p = paragraph._p
    p_pr = p.get_or_add_pPr()
    border = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    border.append(bottom)
    p_pr.append(border)


def p(doc, text="", style=None, align=None):
    para = doc.add_paragraph(style=style)
    if align is not None:
        para.alignment = align
    if text:
        run = para.add_run(text)
        font()(run)
    return para


def add_caption(doc, text):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    font(9, bold=True, color=GRAY)(run)


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, "E8EEF5")
        set_cell_text(cell, h, bold=True, color=NAVY, size=9)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value, size=9)
    if widths:
        for row in table.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return table


def load_font(size=28, bold=False):
    candidates = [
        "/System/Library/Fonts/Supplemental/Arial Bold.ttf" if bold else "/System/Library/Fonts/Supplemental/Arial.ttf",
        "/System/Library/Fonts/Supplemental/Helvetica.ttf",
    ]
    for c in candidates:
        if c and Path(c).exists():
            return ImageFont.truetype(c, size=size)
    return ImageFont.load_default()


def make_logo():
    path = ASSETS / "exosignal_logo.png"
    im = Image.new("RGBA", (900, 260), (255, 255, 255, 0))
    d = ImageDraw.Draw(im)
    d.rounded_rectangle((20, 30, 200, 210), radius=34, fill=(5, 12, 26, 255), outline=(11, 61, 145, 255), width=5)
    d.ellipse((58, 65, 162, 169), outline=(255, 255, 255, 190), width=4)
    d.line((45, 157, 176, 58), fill=(111, 182, 255, 180), width=3)
    d.ellipse((156, 52, 177, 73), fill=(252, 61, 33, 255))
    d.ellipse((43, 174, 66, 197), fill=(111, 182, 255, 255))
    d.text((82, 96), "ES", fill=(255, 255, 255, 255), font=load_font(34, True))
    d.text((230, 64), "EXOSIGNAL", fill=(11, 31, 58, 255), font=load_font(52, True))
    d.text((234, 126), "OBSERVATORY", fill=(11, 61, 145, 255), font=load_font(32, True))
    d.text((236, 172), "Independent transit detection lab", fill=(82, 94, 111, 255), font=load_font(22, False))
    im.save(path)
    return path


def bar_chart(title, data, path, color=(11, 61, 145)):
    w, h = 1200, 640
    im = Image.new("RGB", (w, h), "white")
    d = ImageDraw.Draw(im)
    d.text((50, 32), title, fill=(11, 31, 58), font=load_font(36, True))
    max_v = max(data.values()) or 1
    x0, y0 = 80, 130
    bar_w = 170
    gap = 90
    for i, (label, value) in enumerate(data.items()):
        x = x0 + i * (bar_w + gap)
        bh = int((value / max_v) * 340)
        d.rounded_rectangle((x, y0 + 360 - bh, x + bar_w, y0 + 360), radius=8, fill=color)
        d.text((x, y0 + 380), label.upper(), fill=(31, 41, 55), font=load_font(22, True))
        d.text((x, y0 + 410), f"{value:,}".replace(",", " "), fill=(82, 94, 111), font=load_font(22, False))
    d.line((70, y0 + 360, w - 70, y0 + 360), fill=(210, 218, 230), width=3)
    im.save(path)
    return path


def confusion_chart(matrix, path):
    im = Image.new("RGB", (1000, 720), "white")
    d = ImageDraw.Draw(im)
    d.text((50, 34), "Matrice de confusion — modèle baseline", fill=(11, 31, 58), font=load_font(34, True))
    labels = ["Prédit: non-planète", "Prédit: planète"]
    rows = ["Réel: non-planète", "Réel: planète"]
    max_v = max(max(row) for row in matrix)
    x0, y0 = 300, 170
    cell = 210
    for i, lab in enumerate(labels):
        d.text((x0 + i * cell + 12, y0 - 44), lab, fill=(82, 94, 111), font=load_font(20, True))
    for r in range(2):
        d.text((50, y0 + r * cell + 88), rows[r], fill=(82, 94, 111), font=load_font(20, True))
        for c in range(2):
            val = matrix[r][c]
            intensity = int(235 - 130 * (val / max_v))
            fill = (intensity, intensity + 8, 255)
            d.rectangle((x0 + c * cell, y0 + r * cell, x0 + (c + 1) * cell - 8, y0 + (r + 1) * cell - 8), fill=fill, outline=(11, 31, 58), width=3)
            d.text((x0 + c * cell + 78, y0 + r * cell + 72), str(val), fill=(11, 31, 58), font=load_font(38, True))
    d.text((50, 640), "Lecture : 503 planètes correctement classées, 43 planètes manquées, 116 faux positifs.", fill=(82, 94, 111), font=load_font(20, False))
    im.save(path)
    return path


def pipeline_chart(path):
    im = Image.new("RGB", (1400, 420), "white")
    d = ImageDraw.Draw(im)
    d.text((40, 30), "Pipeline de détection par transit", fill=(11, 31, 58), font=load_font(36, True))
    steps = ["INPUT", "CLEAN", "DETECT", "MEASURE", "SCORE", "VISUALIZE"]
    notes = ["CSV/FITS/MAST", "normalisation", "dips périodiques", "P, depth, SNR", "probabilité", "dashboard"]
    x = 55
    for i, step in enumerate(steps):
        d.rounded_rectangle((x, 140, x + 185, 260), radius=14, fill=(242, 246, 251), outline=(11, 61, 145), width=3)
        d.text((x + 18, 164), f"{i+1:02d}", fill=(252, 61, 33), font=load_font(20, True))
        d.text((x + 18, 192), step, fill=(11, 31, 58), font=load_font(24, True))
        d.text((x + 18, 226), notes[i], fill=(82, 94, 111), font=load_font(18, False))
        if i < len(steps) - 1:
            d.line((x + 195, 200, x + 235, 200), fill=(252, 61, 33), width=4)
            d.polygon([(x + 235, 200), (x + 220, 190), (x + 220, 210)], fill=(252, 61, 33))
        x += 220
    im.save(path)
    return path


def pct(value):
    return f"{value * 100:.2f} %"


def num(value, digits=3):
    return f"{value:.{digits}f}"


def add_chapter_title(doc, title):
    doc.add_page_break()
    for _ in range(5):
        doc.add_paragraph()
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = para.add_run(title)
    font(24, bold=True, color=NAVY, name="Aptos Display")(r)
    add_hline(para, "FC3D21", "10")
    doc.add_page_break()


def add_bullets(doc, items):
    for item in items:
        para = doc.add_paragraph(style="List Bullet")
        run = para.add_run(item)
        font()(run)


def add_numbered(doc, items):
    for item in items:
        para = doc.add_paragraph(style="List Number")
        run = para.add_run(item)
        font()(run)


def add_image_if_exists(doc, path, caption, width=Cm(15.8)):
    if path.exists():
        doc.add_picture(str(path), width=width)
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_caption(doc, caption)


def build_report():
    logo = make_logo()
    class_chart = bar_chart("Répartition des classes", DATASET["eda"]["class_distribution"], ASSETS / "class_balance.png", (11, 61, 145))
    mission_data = {k.upper(): v for k, v in DATASET["eda"]["mission_distribution"].items()}
    mission_chart = bar_chart("Répartition par mission", mission_data, ASSETS / "mission_mix.png", (252, 61, 33))
    confusion = confusion_chart(MODEL["metrics"]["confusion_matrix"]["values"], ASSETS / "confusion_matrix.png")
    pipeline = pipeline_chart(ASSETS / "pipeline.png")

    doc = Document()
    style_document(doc)

    # Cover
    sec = doc.sections[0]
    sec.different_first_page_header_footer = True
    p(doc, "UNIVERSITÉ HASSAN II DE CASABLANCA", align=WD_ALIGN_PARAGRAPH.CENTER).runs[0].bold = True
    p(doc, "FACULTÉ DES SCIENCES BEN M’SICK", align=WD_ALIGN_PARAGRAPH.CENTER)
    p(doc, "DÉPARTEMENT DE MATHÉMATIQUES ET INFORMATIQUE", align=WD_ALIGN_PARAGRAPH.CENTER)
    p(doc, "MASTER BIG DATA & DATA SCIENCE", align=WD_ALIGN_PARAGRAPH.CENTER)
    p(doc, "Année universitaire : 2025 – 2026", align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()
    doc.add_picture(str(logo), width=Cm(13.2))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    title = p(doc, "RAPPORT DE PROJET", align=WD_ALIGN_PARAGRAPH.CENTER)
    title.runs[0].font.size = Pt(20)
    title.runs[0].font.bold = True
    title.runs[0].font.color.rgb = BLUE
    subtitle = p(doc, "ExoSignal Observatory : système intelligent de détection et de priorisation de candidats exoplanètes à partir de courbes de lumière", align=WD_ALIGN_PARAGRAPH.CENTER)
    subtitle.runs[0].font.size = Pt(16)
    subtitle.runs[0].font.bold = True
    subtitle.runs[0].font.color.rgb = NAVY
    doc.add_paragraph()
    add_table(
        doc,
        ["Réalisé par", "Encadré par", "Module / Axe"],
        [["ASHAD Oussama\nEL HADDAJI Hamza", "M. Zakaria EL FAKIR", "Machine Learning, Data Science et analyse de signaux astrophysiques"]],
        [5.4, 4.2, 6.2],
    )
    p(doc, "Casablanca — Juin 2026", align=WD_ALIGN_PARAGRAPH.CENTER)

    add_footer(doc)

    doc.add_page_break()
    p(doc, "DÉDICACES", style="Heading 1")
    p(doc, "Nous dédions ce travail à nos familles, à nos enseignants et à toutes les personnes qui nous ont encouragés pendant la réalisation de ce projet. Leur soutien moral et leur confiance ont accompagné chaque étape, depuis la compréhension scientifique de la photométrie de transit jusqu’à l’intégration d’une application web fonctionnelle.")

    doc.add_page_break()
    p(doc, "REMERCIEMENTS", style="Heading 1")
    p(doc, "Nous exprimons notre profonde gratitude à M. Zakaria EL FAKIR pour ses remarques pédagogiques et son orientation claire : le projet devait dépasser l’exploitation simple d’un catalogue CSV et devenir un système capable d’analyser des courbes de lumière réelles afin de détecter des signaux compatibles avec un transit.")
    p(doc, "Nous remercions également le corps professoral du Master Big Data & Data Science pour la formation scientifique et technique qui a permis de mener ce travail : préparation de données, apprentissage automatique, visualisation, développement web et évaluation rigoureuse des modèles.")

    doc.add_page_break()
    p(doc, "RÉSUMÉ", style="Heading 1")
    p(doc, f"ExoSignal Observatory est une application scientifique interactive destinée à analyser des courbes de lumière Kepler/TESS et à prioriser des candidats exoplanètes. Le système charge une courbe de démonstration, accepte l’import CSV/FITS, nettoie et normalise le signal, détecte les baisses périodiques de flux, extrait des paramètres simples — période, profondeur, durée, SNR, bruit et nombre de dips — puis applique un modèle Machine Learning baseline afin de fournir une probabilité de candidature.")
    p(doc, f"La version actuelle utilise le dataset Hugging Face {DATASET['repo']}, composé de {DATASET['full_dataset_rows']:,} observations couvrant Kepler, K2 et TESS. Le modèle baseline entraîné est un {MODEL['metrics']['model_name']} avec {MODEL['metrics']['features']} caractéristiques. Il atteint F1 = {MODEL['metrics']['f1']:.3f}, ROC-AUC = {MODEL['metrics']['roc_auc']:.3f} et PR-AUC = {MODEL['metrics']['pr_auc']:.3f}. Le dashboard web inclut un atlas interactif d’exoplanètes confirmées, une page planète, un module de prédiction, l’historique des analyses et l’export JSON/PNG.")
    p(doc, "Le modèle ne confirme pas officiellement une exoplanète. Il sert à détecter et prioriser des signaux candidats pour une analyse scientifique plus approfondie.")
    p(doc, "Mots-clés : exoplanète, courbe de lumière, transit, Machine Learning, XGBoost, Random Forest, BLS, TESS, Kepler, FastAPI, React.")

    doc.add_page_break()
    p(doc, "ABSTRACT", style="Heading 1")
    p(doc, "ExoSignal Observatory is an interactive scientific web laboratory for light-curve analysis and exoplanet candidate prioritization. The system loads demo curves, accepts CSV/FITS uploads, cleans and normalizes flux measurements, detects periodic transit-like dips, extracts interpretable parameters, and uses a baseline machine learning model to estimate candidate probability.")
    p(doc, f"The current baseline is an {MODEL['metrics']['model_name']} model trained with {MODEL['metrics']['features']} features. It reaches F1 = {MODEL['metrics']['f1']:.3f}, ROC-AUC = {MODEL['metrics']['roc_auc']:.3f}, and PR-AUC = {MODEL['metrics']['pr_auc']:.3f}. The frontend includes a light-curve canvas, live exoplanet archive view, planet details, prediction panel, history, exports, and command palette. The model is explicitly designed for candidate prioritization, not official discovery confirmation.")

    doc.add_page_break()
    p(doc, "TABLE DES MATIÈRES", style="Heading 1")
    toc = [
        "Dédicaces", "Remerciements", "Résumé", "Abstract", "Liste des figures", "Liste des abréviations",
        "Introduction générale", "Chapitre 1 : Contexte et objectifs", "Chapitre 2 : Machine Learning et transit photométrique",
        "Chapitre 3 : État de l’art", "Chapitre 4 : Méthodologie", "Chapitre 5 : Résultats", "Chapitre 6 : Application web",
        "Conclusion générale", "Bibliographie", "Annexes"
    ]
    add_numbered(doc, toc)

    doc.add_page_break()
    p(doc, "LISTE DES FIGURES", style="Heading 1")
    figures = [
        "Figure 1 : Logo ExoSignal Observatory",
        "Figure 2 : Pipeline fonctionnel de détection par transit",
        "Figure 3 : Page Signal Lab — analyse de courbe de lumière",
        "Figure 4 : Atlas interactif des exoplanètes",
        "Figure 5 : Page Planet — fiche détaillée d’un monde sélectionné",
        "Figure 6 : Signal Vault — historique et actions",
        "Figure 7 : Page Predict — verdict et évidence",
        "Figure 8 : Command Palette — raccourcis et actions",
        "Figure 9 : Répartition des classes du dataset",
        "Figure 10 : Répartition par mission",
        "Figure 11 : Matrice de confusion du modèle baseline",
    ]
    add_bullets(doc, figures)

    doc.add_page_break()
    p(doc, "LISTE DES ABRÉVIATIONS", style="Heading 1")
    add_table(doc, ["Abréviation", "Signification"], [
        ["API", "Application Programming Interface"],
        ["BLS", "Box Least Squares"],
        ["CSV", "Comma-Separated Values"],
        ["EDA", "Exploratory Data Analysis"],
        ["FITS", "Flexible Image Transport System"],
        ["F1-score", "Moyenne harmonique entre précision et rappel"],
        ["KIC", "Kepler Input Catalog"],
        ["MAST", "Mikulski Archive for Space Telescopes"],
        ["ML", "Machine Learning"],
        ["PR-AUC", "Area Under Precision-Recall Curve"],
        ["ROC-AUC", "Area Under Receiver Operating Characteristic Curve"],
        ["SNR", "Signal-to-Noise Ratio"],
        ["TESS", "Transiting Exoplanet Survey Satellite"],
        ["TIC", "TESS Input Catalog"],
        ["XGBoost", "eXtreme Gradient Boosting"],
    ], [3.0, 12.4])

    doc.add_page_break()
    p(doc, "INTRODUCTION GÉNÉRALE", style="Heading 1")
    p(doc, "La recherche d’exoplanètes repose largement sur l’analyse de variations très faibles dans la luminosité des étoiles. Lorsqu’une planète passe devant son étoile, elle provoque une diminution temporaire du flux mesuré : c’est le transit photométrique. Détecter ce phénomène exige une chaîne de traitement capable de distinguer un signal périodique faible du bruit instrumental, de la variabilité stellaire et des faux positifs.")
    p(doc, "Le projet ExoSignal Observatory répond à cet objectif en combinant traitement du signal, Machine Learning et interface web scientifique. L’utilisateur peut charger une courbe de lumière, observer le signal brut et nettoyé, repérer les dips détectés, obtenir les paramètres simples demandés par l’encadrant, puis recevoir un verdict de priorisation sous forme de candidat, signal faible ou non-candidat.")
    p(doc, "La contribution principale du projet est l’intégration complète : backend FastAPI, dataset réel, modèle baseline, visualisation canvas, atlas d’exoplanètes, pages de prédiction et exports. Le système n’est pas un outil officiel de confirmation astronomique ; il constitue une première démonstration fonctionnelle pour filtrer et prioriser les candidats à analyser plus finement.")

    add_chapter_title(doc, "CHAPITRE 1 : CONTEXTE ET OBJECTIFS")
    p(doc, "Introduction", style="Heading 2")
    p(doc, "L’astronomie moderne produit de grandes quantités de données photométriques. Les missions Kepler, K2 et TESS ont permis de collecter des milliers de courbes de lumière, chacune représentant l’évolution temporelle du flux d’une étoile. La détection automatique de transits devient indispensable pour traiter ces volumes.")
    p(doc, "Problématique", style="Heading 2")
    p(doc, "Comment construire un système capable de charger une courbe de lumière réelle ou simulée, la nettoyer, détecter des dips périodiques, extraire des paramètres interprétables et estimer si le signal mérite d’être priorisé comme candidat exoplanète ?")
    p(doc, "Objectifs spécifiques", style="Heading 2")
    add_bullets(doc, [
        "Charger et visualiser des courbes de lumière Kepler/TESS ou des fichiers CSV/FITS.",
        "Nettoyer et normaliser le flux afin de réduire le bruit et la dérive de base.",
        "Détecter les dips périodiques et mesurer période, profondeur, durée, SNR et bruit.",
        "Entraîner et exposer un modèle baseline XGBoost/Random Forest.",
        "Évaluer le modèle avec précision, rappel, F1-score, ROC-AUC, PR-AUC et matrice de confusion.",
        "Développer un dashboard web permettant l’upload, la visualisation, l’analyse et l’export.",
        "Préciser que le résultat est une priorisation de candidat, pas une confirmation officielle.",
    ])

    add_chapter_title(doc, "CHAPITRE 2 : MACHINE LEARNING ET TRANSIT PHOTOMÉTRIQUE")
    p(doc, "Principe du transit", style="Heading 2")
    p(doc, "Une courbe de lumière mesure le flux reçu d’une étoile en fonction du temps. Un transit compatible avec une exoplanète se manifeste par une baisse courte et répétée du flux. La profondeur du dip renseigne approximativement sur le rapport des rayons planète/étoile ; la période correspond à l’intervalle entre deux transits successifs.")
    p(doc, "Paramètres extraits", style="Heading 2")
    add_table(doc, ["Paramètre", "Rôle scientifique", "Exemple démo"], [
        ["Période", "Temps entre deux transits successifs", f"{DEMO['features']['period_estimate']:.3f} jours"],
        ["Profondeur", "Amplitude de la baisse de flux", f"{DEMO['features']['depth_estimate']:.5f}"],
        ["Durée", "Largeur temporelle moyenne du transit", f"{DEMO['features']['duration_estimate']:.3f}"],
        ["SNR", "Force du signal par rapport au bruit", f"{DEMO['features']['snr_estimate']:.2f}"],
        ["Dips", "Nombre de baisses détectées", str(DEMO["features"]["dip_count"])],
        ["Bruit", "Variabilité résiduelle du signal", f"{DEMO['features']['variability']:.5f}"],
    ], [3.0, 8.5, 3.7])
    p(doc, "Modélisation", style="Heading 2")
    p(doc, "Le modèle baseline utilise des caractéristiques numériques extraites des courbes de lumière : statistiques de flux, estimations de périodicité, profondeur, durée, SNR et indicateurs de qualité. L’objectif est de classer les signaux en catégories utiles pour la priorisation.")

    add_chapter_title(doc, "CHAPITRE 3 : ÉTAT DE L’ART")
    p(doc, "Approches classiques", style="Heading 2")
    p(doc, "La recherche de transits utilise traditionnellement des méthodes de traitement du signal comme Box Least Squares (BLS), qui teste une grille de périodes et recherche des baisses de flux cohérentes avec un transit. Ces méthodes sont interprétables et efficaces pour une première analyse.")
    p(doc, "Apprentissage automatique", style="Heading 2")
    p(doc, "Les modèles de type Random Forest et XGBoost sont adaptés à une première baseline car ils traitent bien des caractéristiques tabulaires, résistent au bruit et fournissent une évaluation rapide. Les architectures profondes comme CNN, LSTM et Transformer peuvent ensuite apprendre directement les formes temporelles des courbes.")
    p(doc, "Positionnement du projet", style="Heading 2")
    p(doc, "ExoSignal se situe entre un outil pédagogique et un prototype scientifique : il montre une chaîne complète de bout en bout, depuis la courbe de lumière jusqu’au verdict, tout en gardant la prudence scientifique nécessaire.")

    add_chapter_title(doc, "CHAPITRE 4 : MÉTHODOLOGIE")
    add_image_if_exists(doc, pipeline, "Figure 2 : Pipeline fonctionnel de détection par transit", width=Cm(16.5))
    p(doc, "Source des données", style="Heading 2")
    p(doc, f"Le projet utilise le dataset {DATASET['repo']}. Les fichiers locaux occupent environ {DATASET['files']['train']['size_mb']:.2f} MB pour l’entraînement, {DATASET['files']['val']['size_mb']:.2f} MB pour la validation et {DATASET['files']['test']['size_mb']:.2f} MB pour le test. Le dataset total contient {DATASET['full_dataset_rows']:,} lignes.".replace(",", " "))
    add_table(doc, ["Fichier", "Chemin", "Taille"], [
        ["metadata", DATASET["files"]["metadata"]["path"], f"{DATASET['files']['metadata']['size_mb']:.2f} MB"],
        ["train", DATASET["files"]["train"]["path"], f"{DATASET['files']['train']['size_mb']:.2f} MB"],
        ["val", DATASET["files"]["val"]["path"], f"{DATASET['files']['val']['size_mb']:.2f} MB"],
        ["test", DATASET["files"]["test"]["path"], f"{DATASET['files']['test']['size_mb']:.2f} MB"],
    ], [3.0, 8.5, 3.0])
    add_image_if_exists(doc, class_chart, "Figure 9 : Répartition des classes du dataset", width=Cm(14.8))
    add_image_if_exists(doc, mission_chart, "Figure 10 : Répartition par mission", width=Cm(14.8))
    p(doc, "Nettoyage et détection", style="Heading 2")
    add_numbered(doc, [
        "Lecture du fichier ou de la courbe de démonstration.",
        "Normalisation du flux autour d’une médiane de référence.",
        "Réduction des variations lentes et du bruit local.",
        "Recherche de minima périodiques compatibles avec des transits.",
        "Calcul des métriques simples : période, profondeur, durée, SNR, bruit et nombre de dips.",
        "Envoi des caractéristiques au modèle baseline pour produire la probabilité de candidature.",
    ])

    add_chapter_title(doc, "CHAPITRE 5 : RÉSULTATS")
    p(doc, "Résultat de démonstration", style="Heading 2")
    add_table(doc, ["Mesure", "Valeur"], [
        ["Nombre de points", f"{DEMO['features']['point_count']:,}".replace(",", " ")],
        ["Dips détectés", DEMO["features"]["dip_count"]],
        ["Période estimée", f"{DEMO['features']['period_estimate']:.3f} jours"],
        ["Profondeur estimée", f"{DEMO['features']['depth_estimate']:.5f}"],
        ["SNR", f"{DEMO['features']['snr_estimate']:.2f}"],
        ["Probabilité candidat", pct(DEMO["prediction"]["candidate_probability"])],
    ], [6.5, 7.5])
    add_image_if_exists(doc, SCREENSHOTS["analyze"], "Figure 3 : Page Signal Lab — analyse de courbe de lumière", width=Cm(16.6))
    p(doc, "Performances du modèle", style="Heading 2")
    metrics = MODEL["metrics"]
    add_table(doc, ["Métrique", "Valeur"], [
        ["Modèle", metrics["model_name"]],
        ["Lignes utilisées", f"{metrics['rows_used']:,}".replace(",", " ")],
        ["Nombre de features", metrics["features"]],
        ["Positive rate", pct(metrics["positive_rate"])],
        ["Précision", f"{metrics['precision']:.3f}"],
        ["Rappel", f"{metrics['recall']:.3f}"],
        ["F1-score", f"{metrics['f1']:.3f}"],
        ["ROC-AUC", f"{metrics['roc_auc']:.3f}"],
        ["PR-AUC", f"{metrics['pr_auc']:.3f}"],
        ["Temps d’entraînement", f"{metrics['elapsed_seconds']:.2f} s"],
    ], [6.5, 7.5])
    add_image_if_exists(doc, confusion, "Figure 11 : Matrice de confusion du modèle baseline", width=Cm(14.8))
    p(doc, "Interprétation", style="Heading 2")
    p(doc, "Le F1-score de 0.864 indique un compromis solide entre précision et rappel. Le rappel élevé de 0.921 montre que le modèle rate relativement peu de vrais candidats, ce qui est cohérent avec un outil de priorisation : il vaut mieux signaler certains cas à revoir que manquer des candidats intéressants. Le ROC-AUC de 0.951 indique une bonne capacité de séparation globale.")

    add_chapter_title(doc, "CHAPITRE 6 : APPLICATION WEB")
    p(doc, "Architecture", style="Heading 2")
    add_table(doc, ["Couche", "Technologie", "Rôle"], [
        ["Frontend", "React + Vite + Canvas", "Interface interactive, chart light curve, atlas, pages utilisateur"],
        ["Backend", "FastAPI", "Endpoints santé, demo, analyse, dataset, modèle et entraînement"],
        ["Machine Learning", "XGBoost baseline / Random Forest compatible", "Classification et probabilité candidat"],
        ["Données", "Dataset Hugging Face + Exoplanet Archive + exemples CSV", "Courbes, labels, catalogue et démonstrations"],
    ], [3.2, 4.4, 8.0])
    p(doc, "Pages et fonctionnalités", style="Heading 2")
    add_table(doc, ["Page", "Fonctionnalités principales"], [
        ["Signal Lab", "Run demo, upload CSV/FITS, chart brut/nettoyé, dips, métriques, BLS, MAST, historique, export JSON/PNG"],
        ["Universe", "Atlas interactif, zoom/pan, sélection de monde, fiche de détails, statut archive live/fallback"],
        ["Planet", "Portrait du monde, paramètres orbitaux, prédiction depuis transit projeté, copie des données"],
        ["Signals", "Signal vault, historique des analyses, actions rapides, diagnostics masqués"],
        ["Predict", "Verdict candidat, évidence numérique, upload, demo, prédiction du monde sélectionné, diagnostics modèle"],
    ], [3.2, 12.4])
    add_image_if_exists(doc, SCREENSHOTS["universe"], "Figure 4 : Atlas interactif des exoplanètes", width=Cm(16.6))
    add_image_if_exists(doc, SCREENSHOTS["planet"], "Figure 5 : Page Planet — fiche détaillée d’un monde sélectionné", width=Cm(16.6))
    add_image_if_exists(doc, SCREENSHOTS["signals"], "Figure 6 : Signal Vault — historique et actions", width=Cm(16.6))
    add_image_if_exists(doc, SCREENSHOTS["predict"], "Figure 7 : Page Predict — verdict et évidence", width=Cm(16.6))
    add_image_if_exists(doc, SCREENSHOTS["command"], "Figure 8 : Command Palette — raccourcis et actions", width=Cm(16.6))
    p(doc, "Endpoints backend", style="Heading 2")
    add_table(doc, ["Endpoint", "Méthode", "Rôle"], [
        ["/api/health", "GET", "Vérifie que le backend est disponible"],
        ["/api/demo", "GET", "Retourne une analyse complète sur courbe synthétique"],
        ["/api/analyze", "POST", "Analyse un fichier CSV/FITS envoyé par l’utilisateur"],
        ["/api/dataset", "GET", "Expose les statistiques EDA du dataset"],
        ["/api/model", "GET", "Expose les métriques du modèle baseline"],
        ["/api/train", "POST", "Démarre ou relance un entraînement de démonstration"],
    ], [4.0, 2.5, 9.0])

    doc.add_page_break()
    p(doc, "CONCLUSION GÉNÉRALE", style="Heading 1")
    p(doc, "ExoSignal Observatory démontre une solution fonctionnelle de bout en bout pour l’analyse de courbes de lumière et la priorisation de candidats exoplanètes. Le projet répond aux exigences principales : chargement et visualisation de courbes, nettoyage et normalisation, détection de dips périodiques, extraction de paramètres simples, modèle baseline, évaluation complète et dashboard d’upload.")
    p(doc, "La performance actuelle du modèle est encourageante : F1 = 0.864 et ROC-AUC = 0.951. Ces valeurs ne signifient pas une confirmation scientifique des planètes ; elles indiquent que le modèle est utile pour trier les signaux et orienter une analyse plus avancée.")
    p(doc, "Perspectives", style="Heading 2")
    add_bullets(doc, [
        "Ajouter un modèle CNN 1D pour apprendre directement la forme des transits.",
        "Tester LSTM/Transformer pour mieux exploiter les dépendances temporelles.",
        "Améliorer la validation sur des courbes réelles MAST avec contrôle des faux positifs.",
        "Ajouter des explications de modèle : importance des features, SHAP, score de confiance.",
        "Préparer un mode GPU/cuML ou un déploiement edge pour accélérer l’entraînement.",
    ])

    doc.add_page_break()
    p(doc, "BIBLIOGRAPHIE", style="Heading 1")
    add_bullets(doc, [
        "NASA Exoplanet Archive — catalogue des exoplanètes confirmées et paramètres stellaires.",
        "MAST — archive des produits Kepler, K2 et TESS.",
        "Lightkurve Documentation — analyse des courbes de lumière et recherche BLS.",
        "Hugging Face Dataset : bingbangboom/exoplanet-transit-detection.",
        "Documentation FastAPI, React, Vite et XGBoost.",
    ])

    doc.add_page_break()
    p(doc, "ANNEXES", style="Heading 1")
    p(doc, "A. Exemples de courbes à présenter", style="Heading 2")
    add_table(doc, ["Fichier", "Usage pédagogique"], [
        ["examples/curves/professor_strong_transit_candidate.csv", "Courbe avec dips visibles et verdict candidat"],
        ["examples/curves/professor_weak_noisy_candidate.csv", "Courbe bruitée à discuter avec prudence"],
        ["examples/curves/professor_no_transit_control.csv", "Courbe contrôle sans transit clair"],
    ], [7.2, 8.2])
    p(doc, "B. Checklist des remarques de l’encadrant", style="Heading 2")
    add_bullets(doc, [
        "Courbes de lumière réelles ou simulées visualisées.",
        "Nettoyage et normalisation visibles.",
        "Dips périodiques détectés.",
        "Période, profondeur, durée, SNR et bruit extraits.",
        "Modèle baseline entraîné et métriques affichées.",
        "Évaluation : précision, rappel, F1, ROC-AUC, PR-AUC, matrice de confusion.",
        "Dashboard avec upload, visualisation et probabilité candidat.",
    ])

    doc.save(REPORT)
    return REPORT


if __name__ == "__main__":
    print(build_report())
