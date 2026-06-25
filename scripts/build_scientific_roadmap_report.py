from __future__ import annotations

import json
import os
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
WORKSPACE = ROOT.parents[0]
OUT = ROOT / "deliverables"
ASSETS = ROOT / "deliverables" / "assets"
SCREENSHOTS = WORKSPACE / "outputs"

REPORT_PATH = OUT / "Rapport_ExoSignal_Observatory_Roadmap_Scientifique.docx"


COLORS = {
    "blue": "2E74B5",
    "dark_blue": "1F4D78",
    "ink": "0B2545",
    "muted": "667085",
    "light": "F2F4F7",
    "line": "D9E2EC",
    "green": "146C43",
    "red": "9B1C1C",
    "gold": "7A5A00",
}


def set_cell_text(cell, text: str, bold: bool = False, color: str | None = None) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(str(text))
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_in: float) -> None:
    cell.width = Inches(width_in)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_in * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def make_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, label in enumerate(headers):
        set_cell_text(hdr[i], label, bold=True, color="0B2545")
        shade_cell(hdr[i], "F2F4F7")
        if widths:
            set_cell_width(hdr[i], widths[i])
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
            if widths:
                set_cell_width(cells[i], widths[i])
    doc.add_paragraph()
    return table


def add_bookmark(paragraph, name: str) -> None:
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(abs(hash(name)) % 1000000))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), start.get(qn("w:id")))
    paragraph._p.insert(0, start)
    paragraph._p.append(end)


def add_internal_link(paragraph, text: str, anchor: str) -> None:
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), anchor)
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    style = OxmlElement("w:rStyle")
    style.set(qn("w:val"), "Hyperlink")
    rpr.append(style)
    run.append(rpr)
    t = OxmlElement("w:t")
    t.text = text
    run.append(t)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_external_link(paragraph, text: str, url: str) -> None:
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    style = OxmlElement("w:rStyle")
    style.set(qn("w:val"), "Hyperlink")
    rpr.append(style)
    run.append(rpr)
    t = OxmlElement("w:t")
    t.text = text
    run.append(t)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def heading(doc: Document, text: str, level: int, bookmark: str) -> None:
    p = doc.add_heading(text, level=level)
    add_bookmark(p, bookmark)


def body(doc: Document, text: str, bold_lead: str | None = None) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    if bold_lead:
        lead = p.add_run(bold_lead)
        lead.bold = True
        lead.font.color.rgb = RGBColor.from_string(COLORS["ink"])
        p.add_run(text)
    else:
        p.add_run(text)


def bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.add_run(text)


def add_figure(doc: Document, image_path: Path, caption: str, width: float = 6.2) -> None:
    if not image_path.exists():
        body(doc, f"Figure non disponible: {image_path.name}")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(image_path), width=Inches(width))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(caption)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string(COLORS["muted"])


def page_break(doc: Document) -> None:
    doc.add_page_break()


def configure_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, COLORS["blue"], 16, 8),
        ("Heading 2", 13, COLORS["blue"], 12, 6),
        ("Heading 3", 12, COLORS["dark_blue"], 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("ExoSignal Observatory — Rapport scientifique et feuille de route")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(COLORS["muted"])


def cover(doc: Document) -> None:
    logo = ASSETS / "exosignal_logo.png"
    if logo.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(logo), width=Inches(1.25))
    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("ExoSignal Observatory")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor.from_string(COLORS["ink"])
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("Machine Learning pour la détection et la priorisation de candidats exoplanètes par transit photométrique")
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor.from_string(COLORS["dark_blue"])
    doc.add_paragraph()
    meta = [
        ["Projet", "Système complet d'analyse de courbes de lumière"],
        ["Version du rapport", date.today().isoformat()],
        ["Statut", "Prototype fonctionnel avancé + feuille de route scientifique"],
        ["Positionnement", "Aide à la priorisation; ne confirme pas officiellement une exoplanète"],
    ]
    make_table(doc, ["Champ", "Valeur"], meta, [1.6, 4.7])
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Rapport académique — architecture, résultats, limites et plan vers une utilisation scientifique réelle")
    run.italic = True
    run.font.color.rgb = RGBColor.from_string(COLORS["muted"])
    page_break(doc)


def toc(doc: Document, entries: list[tuple[str, str]]) -> None:
    p = doc.add_heading("Table des matières", level=1)
    add_bookmark(p, "toc")
    for idx, (label, bookmark) in enumerate(entries, start=1):
        row = doc.add_paragraph()
        row.paragraph_format.space_after = Pt(3)
        prefix = row.add_run(f"{idx}. ")
        prefix.bold = True
        add_internal_link(row, label, bookmark)
    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = note.add_run("Les entrées de cette table sont des liens internes cliquables vers les chapitres.")
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor.from_string(COLORS["muted"])
    page_break(doc)


def source_line(doc: Document, label: str, url: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.add_run(f"{label}: ")
    add_external_link(p, url, url)


def build() -> None:
    OUT.mkdir(parents=True, exist_ok=True)
    metrics = json.loads((ROOT / "models" / "metrics.json").read_text())
    rows_used = f"{metrics['rows_used']:,}".replace(",", " ")
    total_rows = "23 567"

    doc = Document()
    configure_doc(doc)
    cover(doc)

    entries = [
        ("Résumé exécutif", "resume"),
        ("Contexte scientifique", "contexte"),
        ("Ce qui a été réalisé", "realise"),
        ("Données et métriques actuelles", "donnees"),
        ("Architecture logicielle", "architecture"),
        ("Visualisations et expérience utilisateur", "visualisations"),
        ("Résultats du modèle baseline", "modele"),
        ("Niveau scientifique visé", "niveau"),
        ("Plan directeur de développement", "plan"),
        ("Besoins matériels et cloud", "materiel"),
        ("Rapport, captures et enregistrements", "captures"),
        ("Limites, risques et formulation correcte", "limites"),
        ("Références", "references"),
    ]
    toc(doc, entries)

    heading(doc, "Résumé exécutif", 1, "resume")
    body(
        doc,
        "ExoSignal Observatory est désormais un prototype fonctionnel avancé: il charge ou simule une courbe de lumière, nettoie le signal, détecte des creux compatibles avec un transit, extrait des paramètres simples et produit une probabilité de candidat. Le système intègre aussi une carte des exoplanètes confirmées de la NASA Exoplanet Archive, une page planète avec NASA Eyes on Exoplanets, une comparaison Terre/Jupiter et une architecture de système stellaire.",
    )
    body(
        doc,
        "le modèle ne confirme pas officiellement une exoplanète. Il priorise des signaux de transit compatibles afin que des scientifiques puissent effectuer une analyse plus approfondie.",
        bold_lead="Point essentiel: ",
    )
    make_table(
        doc,
        ["Indicateur", "Valeur actuelle"],
        [
            ["Lignes de métadonnées", total_rows],
            ["Missions", "Kepler 16 054; TESS 6 028; K2 1 485"],
            ["Dispositions", "PLANET 9 791; NO_SIGNAL 7 840; FALSE_POSITIVE 5 936"],
            ["Modèle baseline", metrics["model_name"]],
            ["Features", str(metrics["features"])],
            ["F1 / ROC-AUC / PR-AUC", f"{metrics['f1']:.3f} / {metrics['roc_auc']:.3f} / {metrics['pr_auc']:.3f}"],
            ["Archive NASA affichée", "6 298 mondes confirmés via TAP dans la requête actuelle"],
        ],
        [2.2, 4.1],
    )

    heading(doc, "Contexte scientifique", 1, "contexte")
    body(
        doc,
        "La méthode du transit photométrique consiste à mesurer la luminosité d'une étoile au cours du temps. Lorsqu'une planète passe devant son étoile depuis notre ligne de visée, la luminosité baisse légèrement. Si ces creux se répètent avec une période stable, une durée cohérente et un rapport signal/bruit suffisant, le signal devient un candidat exoplanète.",
    )
    body(
        doc,
        "Tous les signaux lumineux ne correspondent pas à des exoplanètes: les taches stellaires, les étoiles binaires à éclipses, le bruit instrumental, la contamination par une étoile voisine ou les pulsations stellaires peuvent produire des signaux similaires. C'est pourquoi le rôle du modèle est la priorisation, pas la confirmation définitive.",
    )
    make_table(
        doc,
        ["Source d'ambiguïté", "Pourquoi elle ressemble à un transit", "Contrôle à ajouter"],
        [
            ["Binaire à éclipses", "Deux étoiles se masquent périodiquement", "Test odd/even, profondeur secondaire, centroid shift"],
            ["Taches stellaires", "Rotation de zones sombres sur l'étoile", "Analyse de périodicité et forme non-boxée"],
            ["Bruit instrumental", "Artefacts capteur/cosmic rays", "Nettoyage robuste, masquage outliers, validation multi-secteurs"],
            ["Contamination", "Signal d'une étoile voisine mélangé", "Gaia/centroid, contrôle pixel-level"],
        ],
        [1.6, 2.25, 2.45],
    )

    heading(doc, "Ce qui a été réalisé", 1, "realise")
    bullet(doc, "Frontend React/Vite avec pages Signal Lab, Universe, Planet, Signals et Predict.")
    bullet(doc, "Backend FastAPI avec endpoints santé, dataset, modèle, entraînement, démo et analyse de fichier.")
    bullet(doc, "Analyse de courbe de lumière: visualisation raw/cleaned/dips, score candidat et métriques de transit.")
    bullet(doc, "Upload CSV/FITS et exemples de courbes pour démonstration au professeur.")
    bullet(doc, "Carte interactive NASA Exoplanet Archive avec navigation, zoom, sélection de planète et profil détaillé.")
    bullet(doc, "Intégration NASA Eyes on Exoplanets dans la page Planet, plus liens NASA Science, NASA Archive et Exoplanet.eu.")
    bullet(doc, "Rapport précédent et assets de graphiques: class balance, mission mix, matrice de confusion et pipeline.")
    add_figure(doc, SCREENSHOTS / "exosignal-redesign-analyze.jpg", "Figure 1 — Page Signal Lab: analyse d'une courbe de lumière et résultats.", 6.2)
    add_figure(doc, SCREENSHOTS / "exosignal-redesign-space-planets-visible.jpg", "Figure 2 — Vue Universe: carte interactive des mondes confirmés.", 6.2)

    heading(doc, "Données et métriques actuelles", 1, "donnees")
    make_table(
        doc,
        ["Élément", "Valeur"],
        [
            ["metadata.csv", "33 350 222 octets (~32 MB)"],
            ["train.parquet", "1 528 634 903 octets (~1.4 GB)"],
            ["val.parquet", "191 148 964 octets (~182 MB)"],
            ["test.parquet", "190 954 186 octets (~182 MB)"],
            ["baseline_model.joblib", "405 401 octets"],
            ["feature_columns.json", "1 173 octets"],
        ],
        [2.3, 4.0],
    )
    add_figure(doc, ASSETS / "class_balance.png", "Figure 3 — Répartition des classes dans les données locales.", 5.6)
    add_figure(doc, ASSETS / "mission_mix.png", "Figure 4 — Répartition par mission: Kepler, TESS et K2.", 5.6)

    heading(doc, "Architecture logicielle", 1, "architecture")
    body(
        doc,
        "L'architecture actuelle sépare correctement l'interface utilisateur et les services scientifiques. Le frontend appelle l'API FastAPI; le backend charge les modèles, prépare les features, exécute l'analyse et renvoie un payload unifié pour la visualisation.",
    )
    make_table(
        doc,
        ["Couche", "État actuel", "Évolution nécessaire"],
        [
            ["Frontend", "Très avancé pour une démonstration; NASA map, NASA Eyes, upload, graphiques", "Ajouter Candidate Queue, Evidence Review, recherche TIC/KIC et export scientifique"],
            ["Backend", "Fonctionnel pour démo et analyse locale", "Ajouter jobs asynchrones, stockage résultats, batch scoring, modèle registry"],
            ["Données", "Dataset local déjà volumineux", "Téléchargement MAST automatisé, labels TCE/TOI/KOI, fausses positives riches"],
            ["Modèles", "Baseline ML crédible", "Deep learning 1D CNN, calibration, validation indépendante"],
        ],
        [1.25, 2.55, 2.55],
    )
    add_figure(doc, ASSETS / "pipeline.png", "Figure 5 — Pipeline logique: input, clean, detect, measure, score, evaluate, dashboard.", 6.1)

    heading(doc, "Visualisations et expérience utilisateur", 1, "visualisations")
    body(
        doc,
        "L'expérience visuelle est maintenant un point fort du projet. Elle ne sert pas uniquement la décoration: elle donne un contexte scientifique au signal analysé, permet d'explorer l'archive confirmée, puis de comprendre la planète sélectionnée via son système, son étoile et sa comparaison d'échelle.",
    )
    bullet(doc, "Universe: carte interactive des exoplanètes confirmées avec zoom, déplacement et sélection.")
    bullet(doc, "Planet: viewer NASA Eyes, fiche scientifique, comparaison Terre/Jupiter et système orbital.")
    bullet(doc, "Signal Lab: courbe de lumière interactive, creux détectés, métriques et score.")
    bullet(doc, "Predict: statut du modèle et score de candidat, avec détails masqués par défaut.")
    add_figure(doc, SCREENSHOTS / "exosignal-terminal-space-latest.png", "Figure 6 — Style observatoire/terminal et carte spatiale interactive.", 6.2)

    heading(doc, "Résultats du modèle baseline", 1, "modele")
    body(
        doc,
        "Le modèle actuel est un baseline fort pour un projet académique. Il utilise 58 features et a été entraîné sur 6 000 lignes pour produire une première preuve fonctionnelle. Les performances indiquent que le système capte déjà une partie importante des signaux compatibles avec des transits.",
    )
    make_table(
        doc,
        ["Métrique", "Valeur", "Interprétation"],
        [
            ["Precision", f"{metrics['precision']:.3f}", "Parmi les alertes positives, proportion réellement positive"],
            ["Recall", f"{metrics['recall']:.3f}", "Capacité à retrouver les vrais signaux positifs"],
            ["F1-score", f"{metrics['f1']:.3f}", "Équilibre precision/recall; utile pour classes déséquilibrées"],
            ["ROC-AUC", f"{metrics['roc_auc']:.3f}", "Séparation globale entre classes"],
            ["PR-AUC", f"{metrics['pr_auc']:.3f}", "Qualité sur problème rare/imbalance; plus important que l'accuracy brute"],
            ["Temps entraînement", f"{metrics['elapsed_seconds']:.2f} s", "Entraînement local baseline"],
        ],
        [1.45, 1.0, 3.85],
    )
    cm = metrics["confusion_matrix"]["values"]
    make_table(
        doc,
        ["", "Prédit non-planète", "Prédit planète"],
        [
            ["Vrai non-planète", str(cm[0][0]), str(cm[0][1])],
            ["Vrai planète", str(cm[1][0]), str(cm[1][1])],
        ],
        [1.7, 2.2, 2.2],
    )
    add_figure(doc, ASSETS / "confusion_matrix.png", "Figure 7 — Matrice de confusion du modèle baseline.", 4.8)

    heading(doc, "Niveau scientifique visé", 1, "niveau")
    body(
        doc,
        "L'objectif réaliste n'est pas de remplacer la validation scientifique, mais de produire une liste priorisée de candidats fiables pour suivi. Pour un système réel, l'accuracy seule ne suffit pas: les courbes de lumière sont déséquilibrées, donc F1, recall, precision@K, PR-AUC et calibration sont plus importants.",
    )
    make_table(
        doc,
        ["Niveau", "Objectif métrique", "Signification"],
        [
            ["Démonstration professeur", "F1 ≥ 0.85; ROC-AUC ≥ 0.95; PR-AUC ≥ 0.85", "Prototype convaincant et cohérent avec l'objectif pédagogique"],
            ["Projet académique fort", "F1 ≥ 0.90; ROC-AUC ≥ 0.98; PR-AUC ≥ 0.90", "Résultat robuste sur validation propre Kepler/TESS"],
            ["Triage scientifique", "Recall ≥ 95% sur bons SNR; precision élevée dans le top-K; calibration fiable", "Liste priorisée utile pour analyse humaine et suivi observationnel"],
        ],
        [1.55, 2.55, 2.2],
    )

    heading(doc, "Plan directeur de développement", 1, "plan")
    make_table(
        doc,
        ["Phase", "Objectif", "Travail concret", "Livrable"],
        [
            ["1", "Stabilisation actuelle", "Tests end-to-end, rapport, exemples professeur, nettoyage des textes UI", "Version de soutenance fiable"],
            ["2", "Ingestion MAST réelle", "Recherche TIC/KIC, téléchargement courbes, cache local, parser FITS", "Dataset light curves reproductible"],
            ["3", "Labels scientifiques", "Kepler TCE/KOI, TESS TOI, false positives, eclipsing binaries", "Jeu d'entraînement vérifiable"],
            ["4", "Feature engineering avancé", "BLS power, odd/even, secondary eclipse, centroid/Gaia, robust SNR", "Features scientifiques version 2"],
            ["5", "XGBoost complet", "Entraîner sur dataset élargi, cross-validation, calibration", "Baseline publishable"],
            ["6", "Deep learning", "1D CNN global/local view, puis CNN + features, puis CNN-BiLSTM", "Modèle profond comparable littérature"],
            ["7", "Workflow scientifique", "Batch scoring, Candidate Queue, Evidence Review, export CSV/PDF", "Outil de triage pour analystes"],
            ["8", "Cloud/GPU", "Colab/Kaggle/HF Jobs, tracking expériences, modèle registry", "Entraînement scalable"],
        ],
        [0.55, 1.35, 2.55, 1.75],
    )

    page_break(doc)
    heading(doc, "Besoins matériels et cloud", 1, "materiel")
    body(
        doc,
        "Le PC local suffit pour la démonstration, l'interface, le baseline ML et un petit CNN. Il ne suffit pas pour explorer massivement les millions de courbes TESS/Kepler ou entraîner plusieurs architectures profondes avec validation extensive.",
    )
    make_table(
        doc,
        ["Travail", "PC local", "Cloud/GPU recommandé"],
        [
            ["Frontend + backend + rapport", "Oui", "Non nécessaire"],
            ["XGBoost sur sous-échantillon", "Oui", "Utile mais non obligatoire"],
            ["XGBoost sur dataset complet + CV", "Possible mais lent", "Recommandé"],
            ["1D CNN prototype", "Possible sur petit échantillon", "GPU recommandé"],
            ["CNN/LSTM/Transformer sérieux", "Non recommandé", "GPU nécessaire"],
            ["MAST/TESS à grande échelle", "Stockage et temps limitants", "Cloud storage + jobs asynchrones"],
        ],
        [2.0, 2.1, 2.2],
    )

    heading(doc, "Rapport, captures et enregistrements", 1, "captures")
    body(
        doc,
        "Les captures intégrées ci-dessus documentent l'état visuel et fonctionnel du prototype. Aucun fichier vidéo local n'a été trouvé dans le workspace au moment de la génération du rapport; les enregistrements ci-dessous sont donc listés comme pièces à capturer et joindre séparément.",
    )
    make_table(
        doc,
        ["Enregistrement à joindre", "Durée cible", "Contenu attendu"],
        [
            ["01_signal_lab_demo.mp4", "45-60 s", "Run demo, courbe raw/cleaned, dips et score candidat"],
            ["02_upload_curve.mp4", "45-60 s", "Upload CSV professeur, analyse, métriques et export"],
            ["03_universe_navigation.mp4", "60-90 s", "Zoom/pan sur carte NASA Archive, sélection d'une planète"],
            ["04_planet_profile_nasa_eyes.mp4", "60-90 s", "NASA Eyes embed, comparaison Terre/Jupiter, système orbital"],
            ["05_candidate_queue_future.mp4", "à créer après Phase 7", "Liste priorisée de candidats pour suivi scientifique"],
        ],
        [2.0, 1.0, 3.3],
    )

    heading(doc, "Limites, risques et formulation correcte", 1, "limites")
    bullet(doc, "Le système détecte des signaux compatibles avec un transit; il ne confirme pas officiellement une exoplanète.")
    bullet(doc, "Les exoplanètes déjà présentes dans l'archive NASA sont confirmées; pour elles, le projet montre la visualisation et la logique de détection, pas une redécouverte.")
    bullet(doc, "Les faux positifs astrophysiques restent le risque principal: binaires à éclipses, contamination, activité stellaire et bruit instrumental.")
    bullet(doc, "La métrique accuracy ne doit pas être le critère principal; le rapport doit insister sur recall, precision, F1, PR-AUC, calibration et precision@K.")
    bullet(doc, "Le passage au niveau scientifique nécessite validation externe, jeux de données mieux labellisés et reproductibilité complète.")

    heading(doc, "Références", 1, "references")
    source_line(doc, "NASA — nombre confirmé d'exoplanètes", "https://science.nasa.gov/exoplanets/how-many-exoplanets-are-there/")
    source_line(doc, "NASA Exoplanet Archive — TAP service", "https://exoplanetarchive.ipac.caltech.edu/docs/TAP/usingTAP.html")
    source_line(doc, "TESS data products / MAST", "https://heasarc.gsfc.nasa.gov/docs/tess/data-products.html")
    source_line(doc, "Kepler TCE column definitions", "https://exoplanetarchive.ipac.caltech.edu/docs/API_tce_columns.html")
    source_line(doc, "NASA Eyes on Exoplanets", "https://eyes.nasa.gov/apps/exo/")
    source_line(doc, "Shallue & Vanderburg — Identifying Exoplanets with Deep Learning", "https://arxiv.org/abs/1712.05044")
    source_line(doc, "MNRAS — Accuracy caution for imbalanced exoplanet detection", "https://academic.oup.com/mnras/article/513/4/5505/6472249")

    doc.save(REPORT_PATH)


if __name__ == "__main__":
    build()
    print(REPORT_PATH)
